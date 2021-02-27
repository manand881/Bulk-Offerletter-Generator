import re
from docx import Document

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                print(inline[i].text)
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)


NameRegex = re.compile(r"ReplaceName")
ReplaceNameRegex = r"Anand Mahesh"

PositionRegex = re.compile(r"ReplacePosition")
ReplacePositionRegex = r"God"

DateRegex = re.compile(r"ReplaceDate")
ReplaceDateRegex = r"somedate"

filename = "template.docx"

doc = Document(filename)
docx_replace_regex(doc, NameRegex , ReplaceNameRegex)
docx_replace_regex(doc, PositionRegex , ReplacePositionRegex)
docx_replace_regex(doc, DateRegex , ReplaceDateRegex)

doc.save('template2.docx')