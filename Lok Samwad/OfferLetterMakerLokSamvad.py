import re
import os
import sys
import ctypes
import datetime
import comtypes.client 
from docx import Document
from num2words import num2words

wdFormatPDF = 17

def docx_replace_regex(doc_obj, regex , replace):

    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)

try:
    f=open("Offers.csv","r+")
    Lines=f.readlines()
    Header=Lines[0]
    del Lines[0]
    del Lines[0]
except:
    pass

for line in Lines:

    details=line.split(',')
    Name=details[0]
    Position=details[1]
    DateOfJoining=details[2]
    Term=details[3]
    if "\n" in DateOfJoining:
        DateOfJoining=DateOfJoining[:-1]
    DateOfJoining=datetime.datetime.strptime(DateOfJoining, '%d-%m-%Y').date()
    TempDate=DateOfJoining.strftime("%B")+" "
    if int(DateOfJoining.strftime("%d"))==3:
        TempDate+=str(int(DateOfJoining.strftime("%d")))+"ʳᵈ"
    elif int(DateOfJoining.strftime("%d"))==2:
        TempDate+=str(int(DateOfJoining.strftime("%d")))+"ⁿᵈ"
    elif int(DateOfJoining.strftime("%d"))==1:
        TempDate+=str(int(DateOfJoining.strftime("%d")))+"ˢᵗ"
    elif int(DateOfJoining.strftime("%d"))>3:
        TempDate+=str(int(DateOfJoining.strftime("%d")))+"ᵗʰ"  
    TempDate+=", "+DateOfJoining.strftime("%Y")
    
    NameRegex = re.compile(r"ReplaceName")
    ReplaceNameRegex =  Name 

    PositionRegex = re.compile(r"ReplacePosition")
    ReplacePositionRegex = Position

    DateRegex = re.compile(r"ReplaceDate")
    ReplaceDateRegex = TempDate

    TermRegex = re.compile(r"ReplaceTerm")
    ReplaceTermRegex = "("+Term+")"+" "+num2words(Term)+" months"
    
    try:
        filename = "Lok Samvad Offer Letter.docx"
        doc = Document(filename)
    except:
        MessageBox = ctypes.windll.user32.MessageBoxW
        MessageBox(None, 'Offer Letter Template Not Found', 'Error', 0)
        sys.exit()

    docx_replace_regex(doc, NameRegex , ReplaceNameRegex)
    docx_replace_regex(doc, PositionRegex , ReplacePositionRegex)
    docx_replace_regex(doc, DateRegex , ReplaceDateRegex)
    docx_replace_regex(doc, TermRegex , ReplaceTermRegex)
    OutputFileName=Name+" "+Position+" Offer Letter"
    doc.save(OutputFileName+".docx")
    in_file = os.path.abspath(OutputFileName+".docx")
    out_file = os.path.abspath(OutputFileName+".pdf")
    
    try:
        word = comtypes.client.CreateObject('Word.Application')
    except:
        MessageBox = ctypes.windll.user32.MessageBoxW
        MessageBox(None, 'Microsoft Word Not Installed', 'Error', 0)
        sys.exit()

    word.Visible = True
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()
    os.remove(OutputFileName+".docx")

f.close()
f=open("Offers.csv","w+")
f.writelines(Header)
f.close()